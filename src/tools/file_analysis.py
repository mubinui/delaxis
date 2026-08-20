"""File and image analysis for uploaded documents.

Files arrive through ``POST /api/v1/files`` and land in the uploads directory,
which is also a context root — so :mod:`src.tools.context_tree` can browse them
and these tools can analyse them, with one sandbox shared between both.

Extraction degrades rather than fails. Plain text, CSV, JSON, and image headers
are handled with the standard library alone. PDF, DOCX, and XLSX use optional
packages (``pypdf``, ``python-docx``, ``openpyxl``); when one is missing the
tool returns what it can plus the exact install command, instead of an
ImportError the agent cannot act on.
"""

from __future__ import annotations

import base64
import csv
import json
import mimetypes
import os
import struct
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from src.config.env_compat import env
from src.tools.context_tree import (
    MAX_READ_BYTES,
    ContextAccessError,
    _human_size,
    _is_texty,
    _relative_label,
    resolve_in_sandbox,
)

_PROJECT_ROOT = Path(__file__).resolve().parents[2]

# Formats accepted by the upload endpoint. Anything else is rejected there, so
# the analysers never meet a file type they have no path for.
ALLOWED_SUFFIXES: frozenset[str] = frozenset({
    ".txt", ".md", ".markdown", ".rst", ".log",
    ".csv", ".tsv", ".json", ".jsonl", ".yaml", ".yml", ".xml", ".html",
    ".pdf", ".docx", ".doc", ".xlsx", ".xls", ".pptx",
    ".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".svg", ".tiff",
})

IMAGE_SUFFIXES: frozenset[str] = frozenset({
    ".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".svg", ".tiff",
})

MAX_UPLOAD_BYTES = int(env("DELAXIS_MAX_UPLOAD_BYTES", "26214400") or 26214400)  # 25 MB
MAX_EXTRACT_CHARS = 40_000


def uploads_dir() -> Path:
    """Where uploaded files live. Also a default context root."""
    default = str(_PROJECT_ROOT / "data")
    data_dir = Path(env("DELAXIS_DATA_DIR", default) or default)
    path = data_dir / "uploads"
    path.mkdir(parents=True, exist_ok=True)
    return path


# --------------------------------------------------------------------------- #
# Image headers (no Pillow required)
# --------------------------------------------------------------------------- #


def image_dimensions(path: Path) -> tuple[int, int] | None:
    """Read width/height straight from the file header.

    Pillow would do this too, but it is not a dependency and pulling one in for
    two integers is not worth it. Falls back to Pillow when it happens to be
    installed and the header parse comes up short (TIFF, exotic WebP variants).
    """
    try:
        with path.open("rb") as handle:
            header = handle.read(32)

            if header[:8] == b"\x89PNG\r\n\x1a\n":
                width, height = struct.unpack(">II", header[16:24])
                return int(width), int(height)

            if header[:6] in (b"GIF87a", b"GIF89a"):
                width, height = struct.unpack("<HH", header[6:10])
                return int(width), int(height)

            if header[:2] == b"BM":
                handle.seek(18)
                width, height = struct.unpack("<ii", handle.read(8))
                return abs(int(width)), abs(int(height))

            if header[:4] == b"RIFF" and header[8:12] == b"WEBP":
                chunk = header[12:16]
                if chunk == b"VP8 ":
                    handle.seek(26)
                    width, height = struct.unpack("<HH", handle.read(4))
                    return int(width) & 0x3FFF, int(height) & 0x3FFF
                if chunk == b"VP8L":
                    handle.seek(21)
                    bits = struct.unpack("<I", handle.read(4))[0]
                    return (bits & 0x3FFF) + 1, ((bits >> 14) & 0x3FFF) + 1
                if chunk == b"VP8X":
                    handle.seek(24)
                    raw = handle.read(6)
                    width = 1 + int.from_bytes(raw[0:3], "little")
                    height = 1 + int.from_bytes(raw[3:6], "little")
                    return width, height

            if header[:2] == b"\xff\xd8":
                # JPEG: walk the segment chain to the SOF marker that carries the size.
                handle.seek(2)
                while True:
                    marker = handle.read(2)
                    if len(marker) < 2 or marker[0] != 0xFF:
                        break
                    kind = marker[1]
                    length_bytes = handle.read(2)
                    if len(length_bytes) < 2:
                        break
                    length = struct.unpack(">H", length_bytes)[0]
                    if 0xC0 <= kind <= 0xCF and kind not in (0xC4, 0xC8, 0xCC):
                        payload = handle.read(5)
                        if len(payload) < 5:
                            break
                        height, width = struct.unpack(">HH", payload[1:5])
                        return int(width), int(height)
                    handle.seek(length - 2, os.SEEK_CUR)
    except (OSError, struct.error, ValueError):
        pass

    try:
        from PIL import Image

        with Image.open(path) as image:
            return image.size
    except Exception:
        return None


# --------------------------------------------------------------------------- #
# Text extraction per format
# --------------------------------------------------------------------------- #


def _missing(package: str, install: str) -> dict[str, Any]:
    return {
        "text": "",
        "note": f"Extracting this format needs the optional '{package}' package. Install with: {install}",
        "extractor": "unavailable",
    }


def _extract_pdf(path: Path) -> dict[str, Any]:
    try:
        from pypdf import PdfReader
    except ImportError:
        return _missing("pypdf", "uv pip install pypdf")

    try:
        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages[:100]:
            pages.append(page.extract_text() or "")
        text = "\n\n".join(pages)
        meta = reader.metadata or {}
        return {
            "text": text,
            "extractor": "pypdf",
            "page_count": len(reader.pages),
            "metadata": {
                key.lstrip("/"): str(value)
                for key, value in meta.items()
                if isinstance(key, str)
            },
        }
    except Exception as exc:
        return {"text": "", "extractor": "pypdf", "note": f"PDF could not be parsed: {exc}"}


def _extract_docx(path: Path) -> dict[str, Any]:
    try:
        import docx
    except ImportError:
        return _missing("python-docx", "uv pip install python-docx")

    try:
        document = docx.Document(str(path))
        paragraphs = [para.text for para in document.paragraphs if para.text.strip()]
        tables = []
        for table in document.tables[:20]:
            rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows[:50]]
            tables.append("\n".join(rows))
        body = "\n".join(paragraphs)
        if tables:
            body += "\n\n[tables]\n" + "\n\n".join(tables)
        return {
            "text": body,
            "extractor": "python-docx",
            "paragraph_count": len(paragraphs),
            "table_count": len(document.tables),
        }
    except Exception as exc:
        return {"text": "", "extractor": "python-docx", "note": f"DOCX could not be parsed: {exc}"}


def _extract_xlsx(path: Path) -> dict[str, Any]:
    try:
        import openpyxl
    except ImportError:
        return _missing("openpyxl", "uv pip install openpyxl")

    try:
        workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        parts = []
        sheets = []
        for sheet in workbook.worksheets[:10]:
            sheets.append({"name": sheet.title, "rows": sheet.max_row, "columns": sheet.max_column})
            parts.append(f"# Sheet: {sheet.title}")
            for row in sheet.iter_rows(max_row=200, values_only=True):
                cells = ["" if cell is None else str(cell) for cell in row]
                if any(cell for cell in cells):
                    parts.append(" | ".join(cells))
        workbook.close()
        return {"text": "\n".join(parts), "extractor": "openpyxl", "sheets": sheets}
    except Exception as exc:
        return {"text": "", "extractor": "openpyxl", "note": f"XLSX could not be parsed: {exc}"}


def _extract_csv(path: Path) -> dict[str, Any]:
    try:
        with path.open("r", encoding="utf-8", errors="replace", newline="") as handle:
            sample = handle.read(8192)
            handle.seek(0)
            try:
                dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
            except csv.Error:
                dialect = csv.excel
            reader = csv.reader(handle, dialect)
            rows = []
            for index, row in enumerate(reader):
                if index >= 500:
                    break
                rows.append(row)
    except OSError as exc:
        return {"text": "", "extractor": "csv", "note": f"CSV could not be read: {exc}"}

    if not rows:
        return {"text": "", "extractor": "csv", "row_count": 0}

    header, *body = rows
    # A per-column profile is what makes a CSV actionable — the agent can see
    # which columns are numeric and how sparse they are without reading it all.
    columns = []
    for index, name in enumerate(header):
        values = [row[index] for row in body if index < len(row) and row[index] != ""]
        numeric = [value for value in values if _is_number(value)]
        column: dict[str, Any] = {
            "name": name,
            "filled": len(values),
            "empty": len(body) - len(values),
            "unique": len(set(values)),
            "sample": values[:3],
        }
        if numeric and len(numeric) >= max(1, len(values) // 2):
            numbers = [float(value) for value in numeric]
            column["numeric"] = True
            column["min"] = min(numbers)
            column["max"] = max(numbers)
            column["mean"] = round(sum(numbers) / len(numbers), 4)
        columns.append(column)

    preview = "\n".join(" | ".join(row) for row in rows[:20])
    return {
        "text": preview,
        "extractor": "csv",
        "row_count": len(body),
        "column_count": len(header),
        "columns": columns,
    }


def _is_number(value: str) -> bool:
    try:
        float(value.replace(",", ""))
        return True
    except (ValueError, AttributeError):
        return False


def _json_shape(value: Any, depth: int = 0) -> Any:
    """Describe a JSON document's structure instead of echoing its contents."""
    if depth > 4:
        return "..."
    if isinstance(value, dict):
        return {key: _json_shape(item, depth + 1) for key, item in list(value.items())[:25]}
    if isinstance(value, list):
        if not value:
            return []
        return [_json_shape(value[0], depth + 1), f"... {len(value)} items"] if len(value) > 1 else [
            _json_shape(value[0], depth + 1)
        ]
    return type(value).__name__


def _extract_json(path: Path) -> dict[str, Any]:
    try:
        raw = path.read_text(encoding="utf-8", errors="replace")[:MAX_READ_BYTES]
        parsed = json.loads(raw)
    except (OSError, json.JSONDecodeError) as exc:
        return {"text": "", "extractor": "json", "note": f"JSON could not be parsed: {exc}"}
    return {
        "text": json.dumps(parsed, indent=2)[:MAX_EXTRACT_CHARS],
        "extractor": "json",
        "shape": _json_shape(parsed),
        "top_level_type": type(parsed).__name__,
    }


def _extract_plain(path: Path) -> dict[str, Any]:
    try:
        text = path.read_text(encoding="utf-8", errors="replace")[:MAX_READ_BYTES]
    except OSError as exc:
        return {"text": "", "extractor": "text", "note": f"File could not be read: {exc}"}
    return {"text": text, "extractor": "text", "line_count": text.count("\n") + 1}


def extract_file_text(path: Path) -> dict[str, Any]:
    """Pull text out of ``path`` using the best extractor for its type."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix in (".docx", ".doc"):
        return _extract_docx(path)
    if suffix in (".xlsx", ".xls"):
        return _extract_xlsx(path)
    if suffix in (".csv", ".tsv"):
        return _extract_csv(path)
    if suffix in (".json", ".jsonl"):
        return _extract_json(path)
    if _is_texty(path):
        return _extract_plain(path)
    return {
        "text": "",
        "extractor": "none",
        "note": f"No text extractor for '{suffix}'. Supported: PDF, DOCX, XLSX, CSV, JSON, and plain text.",
    }


# --------------------------------------------------------------------------- #
# Vision
# --------------------------------------------------------------------------- #


def _vision_settings() -> tuple[str, str]:
    provider = env("DELAXIS_VISION_PROVIDER", "") or ""
    model = env("DELAXIS_VISION_MODEL", "") or ""
    return provider, model


def describe_image_with_model(path: Path, question: str) -> dict[str, Any]:
    """Send the image to the configured vision model over the OpenAI-compatible API."""
    provider, model = _vision_settings()
    if not provider or not model:
        return {
            "available": False,
            "note": (
                "No vision model is configured. Set DELAXIS_VISION_PROVIDER and "
                "DELAXIS_VISION_MODEL (e.g. openrouter / openai/gpt-5-mini) to enable "
                "image description. Dimensions and metadata are still reported."
            ),
        }

    try:
        import httpx

        from src.config.provider_registry import resolve_openai_endpoint

        base_url, api_key, auth_required = resolve_openai_endpoint(provider)
    except Exception as exc:
        return {"available": False, "note": f"Vision provider could not be resolved: {exc}"}

    if not base_url:
        return {"available": False, "note": f"Provider '{provider}' has no OpenAI-compatible base_url."}
    if auth_required and not api_key:
        return {"available": False, "note": f"No API key configured for provider '{provider}'."}

    try:
        payload_bytes = path.read_bytes()
    except OSError as exc:
        return {"available": False, "note": f"Image could not be read: {exc}"}

    if len(payload_bytes) > 8_000_000:
        return {
            "available": False,
            "note": f"Image is {_human_size(len(payload_bytes))}; too large to send to a vision model (8MB cap).",
        }

    mime = mimetypes.guess_type(str(path))[0] or "image/png"
    data_uri = f"data:{mime};base64,{base64.b64encode(payload_bytes).decode('ascii')}"
    model_id = model[len(provider) + 1 :] if model.startswith(f"{provider}/") else model

    headers = {"Content-Type": "application/json"}
    if api_key:
        headers["Authorization"] = f"Bearer {api_key}"

    try:
        response = httpx.post(
            f"{base_url.rstrip('/')}/chat/completions",
            headers=headers,
            json={
                "model": model_id,
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": question},
                            {"type": "image_url", "image_url": {"url": data_uri}},
                        ],
                    }
                ],
                "max_tokens": 1000,
            },
            timeout=120,
        )
        if response.status_code >= 400:
            return {
                "available": False,
                "note": f"Vision model returned HTTP {response.status_code}: {response.text[:300]}",
            }
        data = response.json()
        return {
            "available": True,
            "model": model,
            "description": (data["choices"][0]["message"].get("content") or "").strip(),
        }
    except Exception as exc:
        return {"available": False, "note": f"Vision call failed: {exc}"}


# --------------------------------------------------------------------------- #
# Tool entrypoints
# --------------------------------------------------------------------------- #


def list_uploaded_files(pattern: str = "", limit: int = 50) -> str:
    """
    List files that have been uploaded and are available for analysis.

    Args:
        pattern: Optional case-insensitive substring to filter filenames.
        limit: Maximum files to return (1-200, default 50).

    Returns:
        JSON: {"count": int, "files": [{"name", "size", "kind", "uploaded"}]}
    """
    directory = uploads_dir()
    capped = max(1, min(int(limit), 200))
    entries: list[dict[str, Any]] = []

    try:
        candidates = sorted(
            (item for item in directory.iterdir() if item.is_file()),
            key=lambda item: item.stat().st_mtime,
            reverse=True,
        )
    except OSError as exc:
        return json.dumps({"error": f"Uploads directory could not be listed: {exc}"})

    for item in candidates:
        if pattern and pattern.lower() not in item.name.lower():
            continue
        if len(entries) >= capped:
            break
        stat = item.stat()
        entries.append(
            {
                "name": item.name,
                "size": _human_size(stat.st_size),
                "size_bytes": stat.st_size,
                "kind": "image" if item.suffix.lower() in IMAGE_SUFFIXES else "document",
                "uploaded": datetime.fromtimestamp(stat.st_mtime, timezone.utc).isoformat(),
            }
        )

    return json.dumps({"count": len(entries), "files": entries}, indent=2)


def analyze_file(path: str, max_chars: int = 8000) -> str:
    """
    Read an uploaded document and return its text plus a structural summary.

    Handles PDF, DOCX, XLSX, CSV, JSON, and plain text. For a CSV you also get a
    per-column profile (types, ranges, how many values are missing); for JSON,
    the document's shape rather than the whole body.

    Args:
        path: Filename or path within the uploads directory.
        max_chars: Cap on returned text (500-40000, default 8000). The full
            length is always reported so you know what was cut.

    Returns:
        JSON with "text", "extractor", format-specific fields, and file metadata.
    """
    try:
        target = resolve_in_sandbox(path)
    except ContextAccessError as exc:
        return json.dumps({"error": str(exc)})

    if not target.exists():
        return json.dumps({"error": f"'{path}' does not exist. Use list_uploaded_files to see what is available."})
    if target.is_dir():
        return json.dumps({"error": f"'{path}' is a directory. Use context_tree instead."})

    if target.suffix.lower() in IMAGE_SUFFIXES:
        return json.dumps(
            {"error": f"'{path}' is an image. Use analyze_image instead."}
        )

    cap = max(500, min(int(max_chars), MAX_EXTRACT_CHARS))
    result = extract_file_text(target)
    text = result.pop("text", "") or ""
    stat = target.stat()

    return json.dumps(
        {
            "path": _relative_label(target),
            "size": _human_size(stat.st_size),
            "mime_type": mimetypes.guess_type(str(target))[0],
            "text": text[:cap],
            "text_length": len(text),
            "truncated": len(text) > cap,
            **result,
        },
        indent=2,
        default=str,
    )


def analyze_image(path: str, question: str = "Describe this image in detail.") -> str:
    """
    Inspect an uploaded image — dimensions and format always, a written
    description when a vision model is configured.

    Args:
        path: Filename or path of the image within the uploads directory.
        question: What to ask about the image, e.g. "What text appears in this
            screenshot?" or "Is there a signature on this form?".

    Returns:
        JSON with "width", "height", "format", "size", and — when available —
        "description" from the vision model. When no vision model is configured
        the metadata is still returned, with a note explaining how to enable it.
    """
    try:
        target = resolve_in_sandbox(path)
    except ContextAccessError as exc:
        return json.dumps({"error": str(exc)})

    if not target.exists():
        return json.dumps({"error": f"'{path}' does not exist. Use list_uploaded_files to see what is available."})
    if target.suffix.lower() not in IMAGE_SUFFIXES:
        return json.dumps({"error": f"'{path}' is not an image. Use analyze_file instead."})

    stat = target.stat()
    report: dict[str, Any] = {
        "path": _relative_label(target),
        "format": target.suffix.lstrip(".").lower(),
        "size": _human_size(stat.st_size),
        "size_bytes": stat.st_size,
        "mime_type": mimetypes.guess_type(str(target))[0],
    }

    dimensions = image_dimensions(target)
    if dimensions:
        report["width"], report["height"] = dimensions
        report["aspect_ratio"] = round(dimensions[0] / dimensions[1], 3) if dimensions[1] else None

    if target.suffix.lower() == ".svg":
        # SVG is text — its markup is more useful to an agent than a raster
        # description would be.
        try:
            report["svg_source"] = target.read_text(encoding="utf-8", errors="replace")[:4000]
        except OSError:
            pass
        return json.dumps(report, indent=2)

    report["vision"] = describe_image_with_model(target, question)
    return json.dumps(report, indent=2)


def extract_document_text(path: str, max_chars: int = 20000) -> str:
    """
    Return only the plain text of an uploaded document, with no analysis wrapper.

    Use this when you want to feed a document's contents into your own reasoning
    rather than read a summary of it.

    Args:
        path: Filename or path within the uploads directory.
        max_chars: Cap on returned text (500-40000, default 20000).

    Returns:
        JSON: {"path": str, "text": str, "text_length": int, "truncated": bool}
    """
    try:
        target = resolve_in_sandbox(path)
    except ContextAccessError as exc:
        return json.dumps({"error": str(exc)})
    if not target.exists() or target.is_dir():
        return json.dumps({"error": f"'{path}' is not a readable file."})

    cap = max(500, min(int(max_chars), MAX_EXTRACT_CHARS))
    result = extract_file_text(target)
    text = result.get("text", "") or ""
    payload = {
        "path": _relative_label(target),
        "text": text[:cap],
        "text_length": len(text),
        "truncated": len(text) > cap,
        "extractor": result.get("extractor"),
    }
    if result.get("note"):
        payload["note"] = result["note"]
    return json.dumps(payload, indent=2)


def save_uploaded_bytes(filename: str, payload: bytes) -> dict[str, Any]:
    """Persist an upload safely. Used by the API router, not exposed as a tool."""
    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_SUFFIXES:
        raise ValueError(
            f"File type '{suffix or 'unknown'}' is not accepted. "
            f"Allowed: {', '.join(sorted(ALLOWED_SUFFIXES))}"
        )
    if len(payload) > MAX_UPLOAD_BYTES:
        raise ValueError(
            f"File is {_human_size(len(payload))}, over the {_human_size(MAX_UPLOAD_BYTES)} limit."
        )

    # Strip every directory component: an upload named "../../etc/cron.d/x"
    # must not be able to steer where it lands.
    safe_stem = "".join(
        char if char.isalnum() or char in "-_. " else "_" for char in Path(filename).stem
    ).strip() or "upload"
    directory = uploads_dir()
    target = directory / f"{safe_stem}{suffix}"
    counter = 1
    while target.exists():
        target = directory / f"{safe_stem}-{counter}{suffix}"
        counter += 1

    target.write_bytes(payload)
    stat = target.stat()
    return {
        "name": target.name,
        "size": _human_size(stat.st_size),
        "size_bytes": stat.st_size,
        "kind": "image" if suffix in IMAGE_SUFFIXES else "document",
        "mime_type": mimetypes.guess_type(str(target))[0],
        "uploaded": datetime.fromtimestamp(stat.st_mtime, timezone.utc).isoformat(),
    }


def delete_uploaded_file(name: str) -> bool:
    """Remove one uploaded file. Used by the API router, not exposed as a tool."""
    target = uploads_dir() / Path(name).name
    if not target.exists() or not target.is_file():
        return False
    target.unlink()
    return True
