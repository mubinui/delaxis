"""Generate documents agents can hand back to a person.

An agent that produces a report has nowhere to put it: it can return text into a
chat bubble, and that is the end of the workflow. This turns that text into a
file someone can download — a PDF to send on, a Word document to edit, a
spreadsheet to sort.

Content is written as Markdown and rendered into whichever format is asked for,
so one prompt produces every output. Headings, bullets, numbered lists, code
blocks and tables all survive into PDF, Word and HTML; anything else degrades to
a paragraph rather than being dropped.

Formats needing a library (Word, Excel) degrade with a message naming the
package, exactly as the file extractors do. PDF, HTML, Markdown, CSV, JSON and
text need nothing installed.
"""

from __future__ import annotations

import csv
import io
import json
import os
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import structlog

logger = structlog.get_logger(__name__)

_PROJECT_ROOT = Path(__file__).resolve().parents[2]

FORMATS = ("pdf", "docx", "xlsx", "html", "md", "txt", "csv", "json")
MAX_CONTENT_CHARS = 400_000


def generated_dir() -> Path:
    """Where generated documents live, honouring the DELAXIS_DATA_DIR override."""
    default = str(_PROJECT_ROOT / "data")
    path = Path(os.environ.get("DELAXIS_DATA_DIR") or default) / "generated"
    path.mkdir(parents=True, exist_ok=True)
    return path


# --------------------------------------------------------------------------- #
# Markdown, reduced to what a document needs
# --------------------------------------------------------------------------- #


@dataclass(slots=True)
class Block:
    """One structural piece of a document."""

    kind: str          # heading | paragraph | bullet | number | code | table | blank
    text: str = ""
    level: int = 0
    cells: list[str] | None = None


_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET = re.compile(r"^\s*[-*+]\s+(.*)$")
_NUMBER = re.compile(r"^\s*(\d+)[.)]\s+(.*)$")
_TABLE_RULE = re.compile(r"^\s*\|?[\s:|-]+\|[\s:|-]*$")


def parse_markdown(text: str) -> list[Block]:
    """Enough Markdown to lay a document out.

    Not a full parser and not trying to be: the job is to know a heading from a
    paragraph so that a PDF has headings, and to leave everything else legible.
    """
    if not (text or "").strip():
        return []

    blocks: list[Block] = []
    lines = text.replace("\r\n", "\n").split("\n")
    index = 0

    while index < len(lines):
        line = lines[index]

        if line.strip().startswith("```"):
            index += 1
            body: list[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                body.append(lines[index])
                index += 1
            index += 1
            blocks.append(Block("code", "\n".join(body)))
            continue

        if not line.strip():
            blocks.append(Block("blank"))
            index += 1
            continue

        heading = _HEADING.match(line)
        if heading:
            blocks.append(Block("heading", heading.group(2).strip(), level=len(heading.group(1))))
            index += 1
            continue

        # A table is a pipe row followed by a separator row; without the
        # separator it is just a paragraph containing pipes.
        if "|" in line and index + 1 < len(lines) and _TABLE_RULE.match(lines[index + 1]):
            while index < len(lines) and "|" in lines[index]:
                row = lines[index]
                index += 1
                if _TABLE_RULE.match(row):
                    continue
                cells = [cell.strip() for cell in row.strip().strip("|").split("|")]
                blocks.append(Block("table", cells=cells))
            continue

        bullet = _BULLET.match(line)
        if bullet:
            blocks.append(Block("bullet", bullet.group(1).strip()))
            index += 1
            continue

        numbered = _NUMBER.match(line)
        if numbered:
            blocks.append(Block("number", numbered.group(2).strip(), level=int(numbered.group(1))))
            index += 1
            continue

        blocks.append(Block("paragraph", line.strip()))
        index += 1

    return blocks


def _plain(text: str) -> str:
    """Strip inline Markdown for formats that cannot show it."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1 (\2)", text)
    return text


# --------------------------------------------------------------------------- #
# Renderers
# --------------------------------------------------------------------------- #


def _render_pdf(blocks: list[Block], title: str) -> bytes:
    from src.tools.pdf_writer import BOLD, Line, build, wrap

    sizes = {1: 20.0, 2: 15.0, 3: 12.5, 4: 11.5, 5: 11.0, 6: 11.0}
    lines: list[Line] = []

    if title:
        for piece in wrap(title, 22):
            lines.append(Line(piece, size=22, font=BOLD))
        lines.append(Line("", size=8))

    for block in blocks:
        if block.kind == "blank":
            lines.append(Line("", size=5))
        elif block.kind == "heading":
            size = sizes.get(block.level, 11.0)
            for piece in wrap(_plain(block.text), size):
                lines.append(Line(piece, size=size, font=BOLD, space_before=8 if lines else 0))
        elif block.kind in ("bullet", "number"):
            marker = "-  " if block.kind == "bullet" else f"{block.level}. "
            wrapped = wrap(marker + _plain(block.text), 11)
            for position, piece in enumerate(wrapped):
                # Continuation lines are indented under the text, not the marker.
                lines.append(Line(piece if position == 0 else "   " + piece, size=11))
        elif block.kind == "code":
            for row in block.text.split("\n"):
                lines.append(Line("    " + row, size=9.5))
        elif block.kind == "table":
            lines.append(Line("  ".join(_plain(cell) for cell in (block.cells or [])), size=10))
        else:
            for piece in wrap(_plain(block.text), 11):
                lines.append(Line(piece, size=11))

    return build(lines or [Line("")], title=title)


def _render_docx(blocks: list[Block], title: str) -> bytes:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError(
            "Word output needs the 'python-docx' package: uv pip install python-docx"
        ) from None

    document = Document()
    if title:
        document.add_heading(title, level=0)

    table_rows: list[list[str]] = []

    def flush_table() -> None:
        if not table_rows:
            return
        columns = max(len(row) for row in table_rows)
        table = document.add_table(rows=0, cols=columns)
        table.style = "Table Grid"
        for row in table_rows:
            cells = table.add_row().cells
            for position, value in enumerate(row):
                cells[position].text = _plain(value)
        table_rows.clear()

    for block in blocks:
        if block.kind != "table":
            flush_table()
        if block.kind == "heading":
            document.add_heading(_plain(block.text), level=min(block.level, 4))
        elif block.kind == "bullet":
            document.add_paragraph(_plain(block.text), style="List Bullet")
        elif block.kind == "number":
            document.add_paragraph(_plain(block.text), style="List Number")
        elif block.kind == "code":
            document.add_paragraph(block.text, style="Intense Quote")
        elif block.kind == "table":
            table_rows.append(block.cells or [])
        elif block.kind == "paragraph":
            document.add_paragraph(_plain(block.text))
    flush_table()

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _render_xlsx(blocks: list[Block], title: str) -> bytes:
    try:
        from openpyxl import Workbook
    except ImportError:
        raise RuntimeError(
            "Excel output needs the 'openpyxl' package: uv pip install openpyxl"
        ) from None

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = (title or "Sheet1")[:31]

    rows = [block.cells or [] for block in blocks if block.kind == "table"]
    if not rows:
        # No table in the source: one column of the document's lines is still
        # more useful than an empty sheet.
        rows = [[_plain(block.text)] for block in blocks if block.text.strip()]
    for row in rows:
        sheet.append(row)

    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


_HTML_STYLE = """
:root { color-scheme: light dark; }
body { max-width: 46rem; margin: 3rem auto; padding: 0 1.25rem;
       font: 16px/1.65 -apple-system, BlinkMacSystemFont, "Segoe UI", Inter, sans-serif;
       color: #16181d; background: #fff; }
@media (prefers-color-scheme: dark) { body { color: #e6e8ee; background: #14161a; }
       code, pre { background: #1e2128; } th { background: #1e2128; } }
h1, h2, h3 { line-height: 1.25; margin: 2rem 0 .6rem; }
h1 { font-size: 1.9rem; } h2 { font-size: 1.4rem; } h3 { font-size: 1.15rem; }
code, pre { background: #f2f4f7; border-radius: 6px; }
code { padding: .1rem .3rem; font-size: .9em; }
pre { padding: .9rem 1rem; overflow-x: auto; }
table { border-collapse: collapse; width: 100%; margin: 1rem 0; overflow-x: auto; display: block; }
th, td { border: 1px solid #d6dae1; padding: .45rem .7rem; text-align: left; }
th { background: #f2f4f7; font-weight: 600; }
"""


def _inline_html(text: str) -> str:
    """Escape, then restore the inline Markdown that is safe to render."""
    text = (text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
    text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"<em>\1</em>", text)
    text = re.sub(r"`(.+?)`", r"<code>\1</code>", text)
    text = re.sub(r"\[(.+?)\]\((https?://[^\s)]+)\)", r'<a href="\2">\1</a>', text)
    return text


def _render_html(blocks: list[Block], title: str) -> bytes:
    body: list[str] = []
    open_list: str | None = None
    table_rows: list[list[str]] = []

    def close_list() -> None:
        nonlocal open_list
        if open_list:
            body.append(f"</{open_list}>")
            open_list = None

    def flush_table() -> None:
        if not table_rows:
            return
        body.append("<table>")
        for position, row in enumerate(table_rows):
            tag = "th" if position == 0 else "td"
            cells = "".join(f"<{tag}>{_inline_html(cell)}</{tag}>" for cell in row)
            body.append(f"<tr>{cells}</tr>")
        body.append("</table>")
        table_rows.clear()

    for block in blocks:
        if block.kind not in ("bullet", "number"):
            close_list()
        if block.kind != "table":
            flush_table()

        if block.kind == "heading":
            body.append(f"<h{min(block.level, 6)}>{_inline_html(block.text)}</h{min(block.level, 6)}>")
        elif block.kind in ("bullet", "number"):
            wanted = "ul" if block.kind == "bullet" else "ol"
            if open_list != wanted:
                close_list()
                body.append(f"<{wanted}>")
                open_list = wanted
            body.append(f"<li>{_inline_html(block.text)}</li>")
        elif block.kind == "code":
            escaped = block.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            body.append(f"<pre><code>{escaped}</code></pre>")
        elif block.kind == "table":
            table_rows.append(block.cells or [])
        elif block.kind == "paragraph":
            body.append(f"<p>{_inline_html(block.text)}</p>")
    close_list()
    flush_table()

    heading = f"<h1>{_inline_html(title)}</h1>" if title else ""
    page = (
        "<!doctype html>\n<html lang=\"en\">\n<head>\n<meta charset=\"utf-8\">\n"
        f"<meta name=\"viewport\" content=\"width=device-width, initial-scale=1\">\n"
        f"<title>{_inline_html(title or 'Document')}</title>\n<style>{_HTML_STYLE}</style>\n"
        f"</head>\n<body>\n{heading}\n" + "\n".join(body) + "\n</body>\n</html>\n"
    )
    return page.encode("utf-8")


def _render_csv(blocks: list[Block]) -> bytes:
    buffer = io.StringIO()
    writer = csv.writer(buffer)
    rows = [block.cells or [] for block in blocks if block.kind == "table"]
    if not rows:
        rows = [[_plain(block.text)] for block in blocks if block.text.strip()]
    writer.writerows(rows)
    return buffer.getvalue().encode("utf-8")


def _render_json(blocks: list[Block], title: str) -> bytes:
    payload = {
        "title": title,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "blocks": [
            {"kind": block.kind, "level": block.level,
             "text": block.text, "cells": block.cells}
            for block in blocks if block.kind != "blank"
        ],
    }
    return json.dumps(payload, indent=2).encode("utf-8")


# --------------------------------------------------------------------------- #


def _safe_stem(name: str) -> str:
    stem = "".join(
        char if char.isalnum() or char in "-_. " else "_" for char in Path(name).stem
    ).strip()
    return stem or "document"


def generate_document(
    content: str,
    filename: str = "document",
    format: str = "pdf",
    title: str = "",
) -> dict[str, Any]:
    """Write ``content`` out as a downloadable file.

    Args:
        content: The document body, written as Markdown.
        filename: What to call it, without an extension.
        format: One of pdf, docx, xlsx, html, md, txt, csv, json.
        title: Optional heading placed at the top.

    Returns:
        name, format, size, download_url, and error when it failed.
    """
    chosen = (format or "pdf").strip().lower().lstrip(".")
    if chosen not in FORMATS:
        return {"success": False,
                "error": f"unknown format {chosen!r}; expected one of {', '.join(FORMATS)}"}
    if not (content or "").strip():
        return {"success": False, "error": "content is empty"}
    if len(content) > MAX_CONTENT_CHARS:
        return {"success": False,
                "error": f"content is {len(content):,} characters, over the "
                         f"{MAX_CONTENT_CHARS:,} limit"}

    blocks = parse_markdown(content)
    try:
        if chosen == "pdf":
            payload = _render_pdf(blocks, title)
        elif chosen == "docx":
            payload = _render_docx(blocks, title)
        elif chosen == "xlsx":
            payload = _render_xlsx(blocks, title)
        elif chosen == "html":
            payload = _render_html(blocks, title)
        elif chosen == "csv":
            payload = _render_csv(blocks)
        elif chosen == "json":
            payload = _render_json(blocks, title)
        else:  # md, txt — the source is already what was asked for
            header = f"# {title}\n\n" if title and chosen == "md" else (
                f"{title}\n{'=' * len(title)}\n\n" if title else "")
            payload = (header + content).encode("utf-8")
    except RuntimeError as exc:                       # a missing optional package
        return {"success": False, "error": str(exc)}
    except Exception as exc:
        logger.exception("document_generation_failed", format=chosen)
        return {"success": False, "error": f"could not generate the document: {exc}"}

    directory = generated_dir()
    target = directory / f"{_safe_stem(filename)}.{chosen}"
    counter = 1
    while target.exists():
        target = directory / f"{_safe_stem(filename)}-{counter}.{chosen}"
        counter += 1
    target.write_bytes(payload)

    logger.info("document_generated", name=target.name, format=chosen, bytes=len(payload))
    return {
        "success": True,
        "name": target.name,
        "format": chosen,
        "size_bytes": len(payload),
        # Relative, so it works behind whatever host or proxy is serving this.
        "download_url": f"/api/v1/documents/{target.name}",
        "path": str(target),
        "error": None,
    }


def list_generated_documents(limit: int = 100) -> list[dict[str, Any]]:
    """Every document generated so far, newest first."""
    items = []
    for path in generated_dir().iterdir():
        if not path.is_file():
            continue
        stat = path.stat()
        items.append({
            "name": path.name,
            "format": path.suffix.lstrip("."),
            "size_bytes": stat.st_size,
            "created": datetime.fromtimestamp(stat.st_mtime, timezone.utc).isoformat(),
            "download_url": f"/api/v1/documents/{path.name}",
        })
    return sorted(items, key=lambda item: item["created"], reverse=True)[:limit]


def delete_generated_document(name: str) -> bool:
    """Remove one generated document."""
    # Path(name).name strips any directory component, so a crafted name cannot
    # reach outside the generated directory.
    target = generated_dir() / Path(name).name
    if not target.exists() or not target.is_file():
        return False
    target.unlink()
    return True
